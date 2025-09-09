from telegram import Update, ReplyKeyboardMarkup
from telegram.ext import Application, CommandHandler, MessageHandler, filters, ContextTypes
from docx import Document
import os

# =============================
# 🔐 التوكن الخاص بك — تم تعديله ✅
# =============================
BOT_TOKEN = "8104016168:AAGqANlC7WcZJPVf9Ah8efxrcFR3nCAkx3E"

# 💰 معلومات الحساب البنكي
BANK_OWNER = "عمر محمد السهلي"
BANK_NAME = "البنك الراجحي"
BANK_ACCOUNT = "SA00 1234 5678 9012 3456 7890"
PAYMENT_AMOUNT = "50 ريال سعودي"

# تخزين مؤقت للبيانات
user_data = {}

# الأزرار التفاعلية
start_keyboard = [["بدء إنشاء السيرة"]]
start_markup = ReplyKeyboardMarkup(start_keyboard, resize_keyboard=True, one_time_keyboard=True)

payment_keyboard = [["تم الدفع"]]
payment_markup = ReplyKeyboardMarkup(payment_keyboard, resize_keyboard=True, one_time_keyboard=True)

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "👋 أهلاً بك! أنا بوتك الذكي لصنع السيرة الذاتية ✨\n"
        "سأساعدك خطوة بخطوة لصنع سيرة ذاتية احترافية بالإنجليزية — متوافقة مع أنظمة التوظيف.\n\n"
        "اضغط على الزر أدناه للبدء 👇",
        reply_markup=start_markup
    )

async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    text = update.message.text.strip()

    # تهيئة بيانات المستخدم إذا لم تكن موجودة
    if user_id not in user_data:
        user_data[user_id] = {"step": 0}

    step = user_data[user_id]["step"]

    # بدء إنشاء السيرة
    if text == "بدء إنشاء السيرة" and step == 0:
        user_data[user_id]["step"] = 1
        await update.message.reply_text("📝 من فضلك، أدخل اسمك الكامل:")
        return

    # 1. الاسم الكامل
    if step == 1:
        user_data[user_id]["name"] = text
        user_data[user_id]["step"] = 2
        await update.message.reply_text("📱 من فضلك، أدخل رقم جوالك (مثال: +966 50 123 4567):")
        return

    # 2. رقم الجوال
    if step == 2:
        user_data[user_id]["phone"] = text
        user_data[user_id]["step"] = 3
        await update.message.reply_text("📧 من فضلك، أدخل بريدك الإلكتروني:")
        return

    # 3. البريد الإلكتروني
    if step == 3:
        user_data[user_id]["email"] = text
        user_data[user_id]["step"] = 4
        await update.message.reply_text("🎓 من فضلك، أدخل آخر مؤهلك التعليمي (مثال: ثانوية عامة، بكالوريوس...):")
        return

    # 4. الشهادة / التعليم
    if step == 4:
        user_data[user_id]["education"] = text
        user_data[user_id]["step"] = 5
        await update.message.reply_text("💼 أدخل خبرتك العملية الأولى (مثال: منسق مبيعات – شركة المنورة – 2019-الآن)")
        return

    # 5. الخبرات (يمكن إدخال أكثر من خبرة)
    if step == 5:
        if "experiences" not in user_data[user_id]:
            user_data[user_id]["experiences"] = []
        if text.lower() == "تم":
            user_data[user_id]["step"] = 6
            await update.message.reply_text("🛠️ ما هي مهاراتك؟ (أدخل مهارة واحدة في كل مرة)\nمثال: إدارة المشاريع")
            return
        user_data[user_id]["experiences"].append(text)
        await update.message.reply_text("✅ تمت إضافة الخبرة.\nأدخل خبرة أخرى، أو اكتب 'تم' للانتقال للخطوة التالية.")
        return

    # 6. المهارات
    if step == 6:
        if "skills" not in user_data[user_id]:
            user_data[user_id]["skills"] = []
        if text.lower() == "تم":
            user_data[user_id]["step"] = 7
            await update.message.reply_text("🌍 ما هي اللغات التي تتحدثها؟ (أدخل لغة واحدة في كل مرة)\nمثال: العربية (ممتاز)")
            return
        user_data[user_id]["skills"].append(text)
        await update.message.reply_text("✅ تمت إضافة المهارة.\nأدخل مهارة أخرى، أو اكتب 'تم' للانتقال للغات.")
        return

    # 7. اللغات
    if step == 7:
        if "languages" not in user_data[user_id]:
            user_data[user_id]["languages"] = []
        if text.lower() == "تم":
            user_data[user_id]["step"] = 8
            # 🎉 عرض بيانات الحساب البنكي
            await update.message.reply_text(
                f"🎉 تم إنشاء سيرتك الذاتية!\n\n"
                f"💰 من فضلك، حول {PAYMENT_AMOUNT} إلى الحساب التالي:\n\n"
                f"👤 المستفيد: {BANK_OWNER}\n"
                f"🏦 البنك: {BANK_NAME}\n"
                f"💳 رقم الحساب (IBAN): {BANK_ACCOUNT}\n\n"
                f"📲 بعد التحويل، أرسل 'تم الدفع' لتحميل سيرتك الذاتية.",
                reply_markup=payment_markup
            )
            return
        user_data[user_id]["languages"].append(text)
        await update.message.reply_text("✅ تمت إضافة اللغة.\nأدخل لغة أخرى، أو اكتب 'تم' لإنشاء الملف.")
        return

    # 8. التحقق من الدفع
    if step == 8:
        if text == "تم الدفع" or text.lower() == "done payment":
            # 🎁 إنشاء الملف وإرساله
            filename = generate_cv_english(user_id)
            if filename and os.path.exists(filename):
                await update.message.reply_text("✅ تم التحقق من الدفع! إليك سيرتك الذاتية:")
                await context.bot.send_document(chat_id=update.effective_chat.id, document=open(filename, 'rb'))
                await update.message.reply_text("✨ شكرًا لاستخدامك البوت! يمكنك البدء من جديد بأي وقت بإرسال /start")
                # مسح بيانات المستخدم بعد الإرسال
                del user_data[user_id]
            else:
                await update.message.reply_text("❌ حدث خطأ أثناء إنشاء الملف. حاول مرة أخرى لاحقًا.")
            return
        else:
            await update.message.reply_text(
                f"من فضلك، حول {PAYMENT_AMOUNT} إلى:\n"
                f"👤 المستفيد: {BANK_OWNER}\n"
                f"🏦 البنك: {BANK_NAME}\n"
                f"💳 IBAN: {BANK_ACCOUNT}\n"
                f"ثم أرسل 'تم الدفع'.",
                reply_markup=payment_markup
            )

# دالة إنشاء ملف Word بالإنجليزية — متوافق مع الـ ATS
def generate_cv_english(user_id):
    try:
        data = user_data[user_id]
        doc = Document()

        # العنوان الرئيسي
        doc.add_heading(data.get("name", "Unknown Name"), 0)

        # معلومات الاتصال
        contact_info = f"Phone: {data.get('phone', 'N/A')} | Email: {data.get('email', 'N/A')}"
        doc.add_paragraph(contact_info)

        # Professional Summary
        doc.add_heading("Professional Summary", level=1)
        doc.add_paragraph(
            "Results-driven professional with experience in sales, digital solutions, and AI-powered platform development. "
            "Proven track record of generating over 300,000 SAR in annual sales and building interactive e-learning systems. "
            "Seeking to leverage technical and business skills in a dynamic organization."
        )

        # التعليم
        doc.add_heading("Education", level=1)
        doc.add_paragraph(data.get("education", "Not specified"))

        # الخبرات
        doc.add_heading("Work Experience", level=1)
        for exp in data.get("experiences", []):
            doc.add_paragraph(f"• {exp}")

        # المهارات
        doc.add_heading("Skills", level=1)
        for skill in data.get("skills", []):
            doc.add_paragraph(f"• {skill}")

        # اللغات
        doc.add_heading("Languages", level=1)
        for lang in data.get("languages", []):
            doc.add_paragraph(f"• {lang}")

        # اسم الملف
        safe_name = "".join(c for c in data.get("name", "CV") if c.isalnum() or c in (' ', '_', '-')).rstrip()
        filename = f"{safe_name}_ATS_CV.docx"
        doc.save(filename)
        return filename
    except Exception as e:
        print(f"Error generating CV: {e}")
        return None

# تشغيل البوت
if __name__ == '__main__':
    app = Application.builder().token(BOT_TOKEN).build()
    app.add_handler(CommandHandler("start", start))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_message))
    print("✅ Bot is running... Start chatting from Telegram!")
    app.run_polling()
